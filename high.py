import re
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches
import textwrap
import subprocess
import shutil
import requests
from io import BytesIO
OUTPUT_DOCX = Path.home() / "Documents" / f"generated_questions_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
LOCAL_MODEL = "llama3"
CURRICULUM = [
    ("Quantitative Math", "Problem Solving", "Numbers and Operations"),
    ("Quantitative Math", "Problem Solving", "Algebra"),
    ("Quantitative Math", "Geometry and Measurement", "Area & Volume"),
    ("Quantitative Math", "Numbers and Operations", "Fractions, Decimals, & Percents"),
    ("Quantitative Math", "Data Analysis & Probability", "Probability (Basic, Compound Events)"),
    ("Quantitative Math", "Reasoning", "Word Problems"),
]

def detect_options(base_question: str) -> int:
    """Detect how many MCQ options exist in the base question."""
    matches = re.findall(r"\([A-Ea-e]\)", base_question)
    return len(set(matches))
def build_question_format(option_count: int) -> str:
    """Dynamically build question format block based on option count."""
    option_tags = []
    for i in range(option_count):
        if i == option_count - 1:  
            option_tags.append("@@option <Correct Option>")
        else:
            option_tags.append(f"@option <Option {chr(65+i)}>")
    options_block = "\n".join(option_tags)
    return textwrap.dedent(f"""
    @title <Assessment title, meaningful>
    @description <Short assessment description>

    // Use this block for each question when adding Multiple Choice Questions (MCQ)
    @question <Question text here>
    @instruction <Instruction here>
    @difficulty <easy|moderate|hard>
    @Order <Question number>
    {options_block}
    @explanation <Explanation for correct answer>
    @subject <Choose exactly from curriculum list>
    @unit <Choose exactly from curriculum list>
    @topic <Choose exactly from curriculum list>
    @plusmarks 1

    The subject, unit, and topic must come from this curriculum list:
    {chr(10).join(f"- {s} > {u} > {t}" for s, u, t in CURRICULUM)}
    """)
def build_prompt(base_question: str, order_num: int) -> str:
    """Build prompt for LLM with dynamic option formatting."""
    option_count = detect_options(base_question)
    question_format = build_question_format(option_count)
    return f"""
You are an expert educational content creator.
Generate ONE new question that is exactly the same type as the Base Question.
Do not change the mathematical concept — only change surface details like names, objects, and numbers.
Keep LaTeX math notation unchanged.
Preserve any image references in ![](url) format.

Base Question:
{base_question}

Follow this exact output format (tags must be on separate lines):
{question_format}

Ensure:
- Same math skill tested.
- Correct answer logic matches the base question's method.
- Subject, unit, and topic tags match EXACTLY from the curriculum list.
- @subject, @unit, @topic are each on their own line.
"""

def init_app(use_local_model=False):
    global OUTPUT_DOCX
    OUTPUT_DOCX = Path.home() / "Documents" / f"generated_questions_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    print(f"Output folder ready: {OUTPUT_DOCX.parent}")
    print(f"Output file will be: {OUTPUT_DOCX}")
    if use_local_model:
        if shutil.which("ollama") is None:
            raise EnvironmentError("⚠ Ollama not found. Install from https://ollama.ai and run:\nollama pull llama3")
        else:
            print(f"Ollama detected. Using local model: {LOCAL_MODEL}")
    else:
        print("Running in placeholder mode (no AI).")

def call_llm(prompt: str, use_local_model=False) -> str:
    if use_local_model:
        process = subprocess.run(["ollama", "run", LOCAL_MODEL],
                                 input=prompt.encode("utf-8"),
                                 capture_output=True)
        return process.stdout.decode("utf-8", errors="replace").strip()
    else:
        return "@title Placeholder\n@description This is a placeholder question\n..."
def save_to_docx(outputs, path: Path):
    doc = Document()
    doc.add_heading("Generated Questions", level=1)
    for i, out in enumerate(outputs, 1):
        doc.add_paragraph(f"--- Question {i} ---")
        for line in out.splitlines():
            if line.strip().startswith("![](") and line.strip().endswith(")"):
                img_url = line.strip()[4:-1].strip()
                try:
                    if img_url.startswith("http"):
                        response = requests.get(img_url)
                        response.raise_for_status()
                        image_stream = BytesIO(response.content)
                        doc.add_picture(image_stream, width=Inches(4))
                    else:
                        img_path = Path(img_url)
                        if img_path.exists():
                            doc.add_picture(str(img_path), width=Inches(4))
                        else:
                            doc.add_paragraph(f"[Image not found: {img_url}]")
                except Exception as e:
                    doc.add_paragraph(f"[Failed to load image: {img_url}] Error: {e}")
            else:
                doc.add_paragraph(line)
    doc.save(str(path))
    print(f"Saved to {path}")
def main():
    use_local = True
    init_app(use_local_model=use_local)

    print("Enter your base questions. Type 'done' when finished.\n")
    base_questions = []

    while True:
        print("\nPaste base question (including any markdown table). End with empty line:")
        lines = []
        while True:
            line = input()
            if line.strip() == "" and lines:
                break
            elif line.strip().lower() == "done":
                break
            else:
                lines.append(line)
        if not lines:
            if line.strip().lower() == "done":
                break
            else:
                continue
        base_question_text = "\n".join(lines)
        print("Does this question include an image? Enter image URL or local path, or leave blank:")
        image_input = input().strip()
        if image_input:
            base_question_text += f"\n\n![]({image_input})"

        base_questions.append(base_question_text)

        if line.strip().lower() == "done":
            break
    results = []
    for i, bq in enumerate(base_questions, 1):
        prompt = build_prompt(bq, i)
        print(f"\n--- Generating Question {i} ---")
        out = call_llm(prompt, use_local_model=use_local)
        results.append(out)
    save_to_docx(results, OUTPUT_DOCX)
if __name__ == "__main__":
    main()